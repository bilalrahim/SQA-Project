"""
Generate a Word (.docx) comparison report for the individual LLM task.
Usage:
    python scripts/generate_comparison_report_docx.py
"""
import json
from datetime import date
from pathlib import Path
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


BASE = Path(__file__).parent.parent
REQUIREMENTS_FILE   = BASE / "requirements.json"
SELECTED_FILE       = BASE / "selected_rules_individual.json"
MODEL_A_FILE        = BASE / "individual_outputs" / "mistral_test_cases.json"
MODEL_B_FILE        = BASE / "individual_outputs" / "mistral_q2k_test_cases.json"
COMPARISON_FILE     = BASE / "individual_outputs" / "final_comparison_report.json"
OUTPUT_FILE         = BASE / "individual_outputs" / "LLM_Comparison_Report.docx"

MODEL_A_LABEL = "Mistral (Q4_K_M)"
MODEL_B_LABEL = "Mistral Quantized (Q2_K)"

# ---------- helpers ----------

def load(path):
    with open(path) as f:
        return json.load(f)


def set_cell_bg(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


def bold_run(para, text, size=None, color=None):
    run = para.add_run(text)
    run.bold = True
    if size:
        run.font.size = Pt(size)
    if color:
        run.font.color.rgb = RGBColor(*color)
    return run


def heading(doc, text, level):
    p = doc.add_heading(text, level=level)
    return p


def add_labeled_para(doc, label, value):
    p = doc.add_paragraph()
    bold_run(p, f"{label}: ")
    p.add_run(str(value))
    return p


def add_kv_table(doc, rows, col_widths=(2.2, 4.0)):
    table = doc.add_table(rows=len(rows), cols=2)
    table.style = "Table Grid"
    for i, (key, val) in enumerate(rows):
        cells = table.rows[i].cells
        cells[0].width = Inches(col_widths[0])
        cells[1].width = Inches(col_widths[1])
        kp = cells[0].paragraphs[0]
        bold_run(kp, key)
        cells[1].paragraphs[0].add_run(str(val))
    return table


def tick(val):
    return "Yes" if val else "No"


def steps_text(steps):
    if not steps:
        return "—"
    if isinstance(steps, list):
        parts = []
        for s in steps:
            if isinstance(s, dict):
                num = s.get("step_number", "")
                desc = s.get("step_description") or s.get("action") or s.get("description", "")
                parts.append(f"{num}. {desc}".strip(". "))
            else:
                parts.append(str(s))
        return "\n".join(parts)
    return str(steps)


# ---------- main ----------

def build():
    requirements  = {r["requirement_id"]: r for r in load(REQUIREMENTS_FILE)}
    selected_ids  = load(SELECTED_FILE)
    model_a_cases = {c["requirement_id"]: c for c in load(MODEL_A_FILE)}
    model_b_cases = {c["requirement_id"]: c for c in load(MODEL_B_FILE)}
    comparison    = load(COMPARISON_FILE)
    summary       = comparison["summary"]
    per_req       = {r["requirement_id"]: r for r in comparison["per_requirement"]}

    doc = Document()

    # ---------- margins ----------
    for section in doc.sections:
        section.top_margin    = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin   = Inches(1.2)
        section.right_margin  = Inches(1.2)

    # ===== TITLE PAGE =====
    title_para = doc.add_paragraph()
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    bold_run(title_para, "LLM-Based Test Case Generation\nComparison Report", size=18)

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub.add_run(f"21 CFR 117.130 – Hazard Analysis Requirements\n"
                f"Individual Task | SQA 2026 – Auburn University\n"
                f"Date: {date.today().strftime('%B %d, %Y')}")

    doc.add_paragraph()

    # ===== SECTION 1: Overview =====
    heading(doc, "1. Overview", 1)
    doc.add_paragraph(
        "This report documents the individual LLM-based test case generation task for "
        "SQA 2026 at Auburn University. Two Mistral language model variants were used "
        "to independently generate test cases for five atomic requirements extracted from "
        "21 CFR §117.130 (Hazard Analysis). The outputs are compared across three "
        "dimensions: Coverage, Correctness, and Completeness."
    )

    # ===== SECTION 2: Models =====
    heading(doc, "2. Models Used", 1)
    add_kv_table(doc, [
        ("Model A", MODEL_A_LABEL),
        ("Model B", MODEL_B_LABEL),
        ("Runtime", "Ollama (local)"),
        ("Base architecture", "Mistral 7B Instruct"),
        ("Key difference",
         "Q4_K_M uses 4-bit quantization with medium precision. "
         "Q2_K uses 2-bit quantization — smaller, faster, but lower fidelity."),
    ])

    doc.add_paragraph()

    # ===== SECTION 3: Selected Requirements =====
    heading(doc, "3. Selected Requirements", 1)
    doc.add_paragraph(
        "Five atomic requirements were selected from the parsed requirements.json file "
        "spanning all three CFR sub-sections (a), (b), and (c)."
    )

    req_table = doc.add_table(rows=1, cols=3)
    req_table.style = "Table Grid"
    hdr = req_table.rows[0].cells
    for cell, label in zip(hdr, ["Requirement ID", "Description", "Source"]):
        set_cell_bg(cell, "1F4E79")
        p = cell.paragraphs[0]
        bold_run(p, label, color=(255, 255, 255))

    for rid in selected_ids:
        req = requirements.get(rid, {})
        row = req_table.add_row().cells
        row[0].paragraphs[0].add_run(rid)
        row[1].paragraphs[0].add_run(req.get("description", "—"))
        row[2].paragraphs[0].add_run(req.get("source", "—"))

    doc.add_paragraph()

    # ===== SECTION 4: Generated Test Cases per Model =====
    heading(doc, "4. Generated Test Cases", 1)

    for model_label, cases_map in [(MODEL_A_LABEL, model_a_cases), (MODEL_B_LABEL, model_b_cases)]:
        heading(doc, f"4.{'1' if 'Q4' in model_label else '2'}  {model_label}", 2)
        for rid in selected_ids:
            tc = cases_map.get(rid)
            if not tc:
                doc.add_paragraph(f"{rid}: No test case generated.")
                continue

            heading(doc, f"{tc['test_case_id']} – {rid}", 3)
            add_kv_table(doc, [
                ("Test Case ID",     tc.get("test_case_id", "—")),
                ("Requirement ID",   tc.get("requirement_id", "—")),
                ("Description",      tc.get("description", "—")),
                ("Input Data",       str(tc.get("input_data", "—"))),
                ("Expected Output",  str(tc.get("expected_output", "—"))),
                ("Steps",            steps_text(tc.get("steps", []))),
                ("Notes",            tc.get("notes") or "—"),
            ])
            doc.add_paragraph()

    # ===== SECTION 5: Comparison =====
    heading(doc, "5. Comparison Analysis", 1)

    # 5.1 Summary table
    heading(doc, "5.1  Summary", 2)
    add_kv_table(doc, [
        ("Requirements selected",               summary["selected_requirements"]),
        ("Requirements covered by both models",  summary["requirements_covered_by_both_models"]),
        ("Full coverage achieved",               tick(summary["full_coverage"])),
        ("All required fields present",          tick(summary["all_required_fields_present"])),
        ("Correctness requires manual review",   tick(summary["correctness_requires_manual_review"])),
    ])
    doc.add_paragraph()

    # 5.2 Per-requirement breakdown
    heading(doc, "5.2  Per-Requirement Breakdown", 2)

    comparison_table = doc.add_table(rows=1, cols=7)
    comparison_table.style = "Table Grid"
    col_headers = [
        "Req ID",
        f"Coverage\n{MODEL_A_LABEL}",
        f"Coverage\n{MODEL_B_LABEL}",
        f"Complete\n{MODEL_A_LABEL}",
        f"Complete\n{MODEL_B_LABEL}",
        f"Heuristic\nCorrectness\n{MODEL_A_LABEL}",
        f"Heuristic\nCorrectness\n{MODEL_B_LABEL}",
    ]
    hdr_cells = comparison_table.rows[0].cells
    for cell, label in zip(hdr_cells, col_headers):
        set_cell_bg(cell, "1F4E79")
        p = cell.paragraphs[0]
        bold_run(p, label, color=(255, 255, 255))

    for rid in selected_ids:
        r = per_req.get(rid, {})
        cov  = r.get("coverage", {})
        comp = r.get("completeness", {})
        corr = r.get("correctness", {})

        row = comparison_table.add_row().cells
        row[0].paragraphs[0].add_run(rid)
        row[1].paragraphs[0].add_run("Yes" if cov.get("mistral_Q4_K_M") else "No")
        row[2].paragraphs[0].add_run("Yes" if cov.get("mistral_Q2_K") else "No")
        row[3].paragraphs[0].add_run(tick(comp.get("mistral_Q4_K_M", {}).get("has_all_required_fields", False)))
        row[4].paragraphs[0].add_run(tick(comp.get("mistral_Q2_K", {}).get("has_all_required_fields", False)))
        row[5].paragraphs[0].add_run(tick(corr.get("mistral_Q4_K_M", {}).get("heuristic_match", False)))
        row[6].paragraphs[0].add_run(tick(corr.get("mistral_Q2_K", {}).get("heuristic_match", False)))

    doc.add_paragraph()

    # 5.3 Detailed correctness discussion
    heading(doc, "5.3  Correctness Discussion", 2)
    correctness_rows = []
    for rid in selected_ids:
        r = per_req.get(rid, {})
        corr = r.get("correctness", {})
        req_desc = r.get("requirement_description", "—")
        a_terms = ", ".join(corr.get("mistral_Q4_K_M", {}).get("overlap_terms", [])) or "none"
        b_terms = ", ".join(corr.get("mistral_Q2_K",   {}).get("overlap_terms", [])) or "none"
        a_match = tick(corr.get("mistral_Q4_K_M", {}).get("heuristic_match", False))
        b_match = tick(corr.get("mistral_Q2_K",   {}).get("heuristic_match", False))
        correctness_rows.append(
            f"{rid} ({req_desc})\n"
            f"  {MODEL_A_LABEL}: heuristic match = {a_match}, shared terms: {a_terms}\n"
            f"  {MODEL_B_LABEL}: heuristic match = {b_match}, shared terms: {b_terms}"
        )
    for block in correctness_rows:
        doc.add_paragraph(block, style="List Bullet")

    doc.add_paragraph(
        "Note: Heuristic correctness is based on keyword overlap between the test case "
        "description and the requirement text. Manual review is required to confirm "
        "that each test case actually validates the intent of its requirement."
    )

    # ===== SECTION 6: Analysis & Conclusions =====
    heading(doc, "6. Analysis and Conclusions", 1)

    heading(doc, "6.1  Coverage", 2)
    doc.add_paragraph(
        f"Both {MODEL_A_LABEL} and {MODEL_B_LABEL} produced at least one test case for "
        f"each of the five selected requirements, achieving 100% coverage. "
        "Neither model skipped a requirement."
    )

    heading(doc, "6.2  Completeness", 2)
    doc.add_paragraph(
        f"{MODEL_A_LABEL} produced well-populated test cases with specific input data, "
        "detailed expected outputs, and multi-step procedures for all 5 requirements. "
        "Optional fields (steps and notes) were consistently used and meaningful."
    )
    doc.add_paragraph(
        f"{MODEL_B_LABEL} (heavier quantization) produced structurally valid test cases "
        "but left input_data and expected_output fields empty for 3 out of 5 requirements, "
        "receiving placeholder values during post-processing. Steps were structured as "
        "objects rather than plain strings in some cases, indicating reduced instruction-following "
        "fidelity under Q2_K quantization."
    )

    heading(doc, "6.3  Correctness", 2)
    doc.add_paragraph(
        f"Both models performed similarly on the two semantically rich requirements "
        f"(REQ-117.130-002A: Biological hazards and REQ-117.130-003A1: severity/probability), "
        "achieving a heuristic match on both. For shorter requirements such as "
        "REQ-117.130-001A (Conduct hazard analysis) and REQ-117.130-001F "
        "(Hazard analysis must be written), neither model achieved the heuristic threshold, "
        "though both descriptions are directionally relevant. Manual review is recommended."
    )

    heading(doc, "6.4  Overall Verdict", 2)
    doc.add_paragraph(
        f"{MODEL_A_LABEL} is the stronger model for this task. It produced richer, "
        "more domain-specific test cases with fully populated required and optional fields. "
        f"{MODEL_B_LABEL} is usable but shows clear degradation in completeness due to "
        "aggressive quantization, particularly in generating concrete input data and "
        "expected outputs. For regulatory test case generation (CFR compliance), "
        "the higher-precision model is strongly preferred."
    )

    # ===== SECTION 7: Deliverables =====
    heading(doc, "7. Deliverables", 1)
    add_kv_table(doc, [
        ("requirements.json",              "26 parsed atomic rules from 21 CFR 117.130"),
        ("selected_rules_individual.json", "5 requirements selected for individual task"),
        ("mistral_test_cases.json",        f"5 test cases generated by {MODEL_A_LABEL}"),
        ("mistral_q2k_test_cases.json",    f"5 test cases generated by {MODEL_B_LABEL}"),
        ("final_comparison_report.json",   "Machine-readable comparison report"),
        ("LLM_Comparison_Report.docx",     "This report"),
        ("generate_llm_test_cases.py",     "Script to generate LLM test cases via Ollama"),
        ("compare_llm_outputs.py",         "Script to compare two model output files"),
    ])

    # ===== save =====
    OUTPUT_FILE.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUTPUT_FILE)
    print(f"Saved -> {OUTPUT_FILE}")


if __name__ == "__main__":
    build()
