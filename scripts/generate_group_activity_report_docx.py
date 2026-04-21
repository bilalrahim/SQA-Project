import json
from datetime import date
from pathlib import Path

from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH


ROOT = Path(__file__).resolve().parent.parent
VV_REPORT = ROOT / "vv_report.json"
FORENSIC_LOG = ROOT / "forensic.log"
WORKFLOW_FILE = ROOT / ".github" / "workflows" / "vv_pipeline.yml"
OUTPUT_FILE = ROOT / "individual_outputs" / "Group_Activities_and_Learnings_Report.docx"


def load_json(path: Path):
    with path.open("r", encoding="utf-8") as handle:
        return json.load(handle)


def add_title(doc: Document):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Project Activities and Lessons Learned Report")
    run.bold = True
    run.font.size = Pt(18)

    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p2.add_run("Course: SQA 2026 | Project: 21 CFR Atomic Rules")

    p3 = doc.add_paragraph()
    p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p3.add_run(f"Date: {date.today().strftime('%B %d, %Y')}")

    doc.add_paragraph()


def add_overview(doc: Document):
    doc.add_heading("1. Project Overview", level=1)
    doc.add_paragraph(
        "This report summarizes the activities performed in the group project and the "
        "key lessons learned while implementing requirement extraction, test case "
        "generation, verification, validation, forensic logging, and CI automation."
    )


def add_activities(doc: Document, vv_data: dict):
    summary = vv_data.get("summary", {})
    selected = vv_data.get("verification", {}).get("selected_requirements", {})
    integrity = vv_data.get("verification", {}).get("test_case_integrity", {})

    doc.add_heading("2. Activities Performed", level=1)

    items = [
        "Task 0: Repository setup and team collaboration structure were established.",
        "Task 1: Parsed CFR markdown into structured requirement entries (requirements.json).",
        "Task 1: Selected 10 atomic rules and generated expected_structure.json mapping parent IDs to child letters.",
        "Task 2: Generated minimal test cases (test_cases.json), one test case per selected requirement.",
        "Task 3: Executed automated verification and validation scripts.",
        "Task 4: Integrated forensic logging across verification/validation workflow.",
        "Configured GitHub Actions pipeline (.github/workflows/vv_pipeline.yml) to run generation, V&V checks, and upload forensic artifacts.",
        "Completed individual LLM task with two Mistral variants and generated comparison artifacts.",
    ]

    for item in items:
        doc.add_paragraph(item, style="List Bullet")

    doc.add_paragraph()
    doc.add_heading("2.1 Measured Outcomes", level=2)
    doc.add_paragraph(f"Verification Passed: {summary.get('verification_passed', False)}")
    doc.add_paragraph(f"Validation Passed: {summary.get('validation_passed', False)}")
    doc.add_paragraph(f"Overall Passed: {summary.get('overall_passed', False)}")
    doc.add_paragraph(
        f"Selected Requirement Count: {selected.get('selected_requirement_count', 'N/A')} "
        f"(Expected: {selected.get('expected_selected_requirement_count', 'N/A')})"
    )
    doc.add_paragraph(
        f"Generated Test Case Count: {integrity.get('generated_test_case_count', 'N/A')} "
        f"(Expected: {integrity.get('expected_test_case_count', 'N/A')})"
    )


def add_learnings(doc: Document):
    doc.add_heading("3. What We Learned", level=1)

    lessons = [
        "Regulatory text must be normalized before automation; numbering patterns can break naive parsers.",
        "Requirement IDs and parent-child mappings are critical for traceability and reliable downstream validation.",
        "Minimal test cases can still be effective if they strictly map to requirement intent and expected output.",
        "Forensic logging significantly improves auditability by recording every verification and validation action.",
        "CI workflows reduce manual mistakes by automatically regenerating and rechecking artifacts on each push/PR.",
        "LLM outputs require post-processing and consistency checks, especially with heavily quantized models.",
        "Comparing model variants highlights tradeoffs between output quality and inference efficiency.",
    ]

    for lesson in lessons:
        doc.add_paragraph(lesson, style="List Bullet")


def add_artifacts(doc: Document):
    doc.add_heading("4. Key Artifacts Produced", level=1)
    artifacts = [
        "requirements.json",
        "expected_structure.json",
        "test_cases.json",
        "forensic.log",
        "vv_report.json",
        ".github/workflows/vv_pipeline.yml",
        "individual_outputs/mistral_test_cases.json",
        "individual_outputs/mistral_q2k_test_cases.json",
        "individual_outputs/final_comparison_report.json",
        "individual_outputs/LLM_Comparison_Report.docx",
    ]
    for artifact in artifacts:
        doc.add_paragraph(artifact, style="List Bullet")


def add_screenshot_placeholders(doc: Document):
    doc.add_heading("5. Screenshot Placeholders", level=1)
    doc.add_paragraph("Insert screenshots below before final submission.")

    placeholders = [
        ("Screenshot 1: GitHub Actions successful run", "Paste screenshot of the green V&V Pipeline run from the GitHub Actions tab."),
        ("Screenshot 2: Forensic log excerpt", "Paste screenshot showing forensic.log execution entries (verification + validation)."),
        ("Screenshot 3: Uploaded CI artifacts", "Paste screenshot showing uploaded artifacts: forensic-log and vv-report."),
        ("Screenshot 4: Verification/Validation report evidence", "Paste screenshot of vv_report.json summary showing all checks passed."),
    ]

    for title, instruction in placeholders:
        doc.add_heading(title, level=2)
        doc.add_paragraph(instruction)
        box = doc.add_table(rows=1, cols=1)
        box.style = "Table Grid"
        cell = box.cell(0, 0)
        cell.text = "[Insert Screenshot Here]"
        cell.width = Inches(6)
        doc.add_paragraph()


def main():
    vv_data = load_json(VV_REPORT)

    doc = Document()
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    add_title(doc)
    add_overview(doc)
    add_activities(doc, vv_data)
    add_learnings(doc)
    add_artifacts(doc)
    add_screenshot_placeholders(doc)

    OUTPUT_FILE.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUTPUT_FILE)
    print(f"Saved report: {OUTPUT_FILE}")


if __name__ == "__main__":
    main()
